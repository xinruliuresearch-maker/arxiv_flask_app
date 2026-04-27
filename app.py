# app.py
import html
import os
import re
import shutil
import tempfile
import threading
import time
import uuid
from collections import defaultdict
from dataclasses import dataclass, field
from typing import Any, Dict, Iterable, List, Optional, Set, Tuple

import feedparser
import fitz  # PyMuPDF
import requests
from flask import Flask, jsonify, render_template, request, send_file, url_for
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


app = Flask(__name__)
app.json.ensure_ascii = False

GENERATED_FILES: Dict[str, Dict[str, Any]] = {}
GENERATED_FILES_LOCK = threading.Lock()
FILE_TTL_SECONDS = 60 * 60

TITLE_FONTS = ("Cambria", "黑体")
BODY_FONTS = ("Times New Roman", "宋体")
HEADING_FONTS = ("Arial", "黑体")
CAPTION_FONTS = ("Times New Roman", "宋体")

FIGURE_CAPTION_RE = re.compile(r"^(fig(?:ure)?\.?\s*\d+[:.\-]?\s*.*|图\s*\d+[：:.\-]?\s*.*)$", re.I)
KEYWORDS_RE = re.compile(r"^(keywords?|index terms?|关键词)\s*[:：-]\s*(.+)$", re.I)
ABSTRACT_RE = re.compile(r"^(abstract|摘要)\s*[:：]?\s*(.*)$", re.I)
REFERENCE_HEADING_RE = re.compile(r"^(references|bibliography|参考文献)$", re.I)
ARXIV_NOISE_RE = re.compile(r"\barxiv:\s*\d{4}\.\d{4,5}", re.I)

OLLAMA_API_URL = os.environ.get("OLLAMA_API_URL", "http://127.0.0.1:11434/api/generate")
OLLAMA_MODEL = os.environ.get("OLLAMA_MODEL", "deepseek-r1:8b")
TRANSLATION_SEPARATOR = "\n<<<SEG>>>\n"


@dataclass
class TextBlock:
    id: str
    page_number: int
    block_number: int
    bbox: Tuple[float, float, float, float]
    text: str
    font_size: float
    bold: bool
    line_count: int

    @property
    def x0(self) -> float:
        return float(self.bbox[0])

    @property
    def y0(self) -> float:
        return float(self.bbox[1])

    @property
    def x1(self) -> float:
        return float(self.bbox[2])

    @property
    def y1(self) -> float:
        return float(self.bbox[3])


@dataclass
class FigureItem:
    page_number: int
    order_y: float
    image_path: str
    caption: str
    width_ratio: float
    is_wide: bool


@dataclass
class TableItem:
    caption: str
    rows: List[List[str]]
    page_number: int
    order_y: float


@dataclass
class SectionItem:
    type: str
    text: str = ""
    level: int = 0
    figure: Optional[FigureItem] = None
    table: Optional[TableItem] = None


@dataclass
class SectionBlock:
    heading: str
    level: int
    items: List[SectionItem] = field(default_factory=list)


@dataclass
class ArticleStructure:
    title: str
    title_cn: str = ""
    title_en: str = ""
    authors: List[str] = field(default_factory=list)
    affiliations: List[str] = field(default_factory=list)
    funding: List[str] = field(default_factory=list)
    author_notes: List[str] = field(default_factory=list)
    abstract: str = ""
    abstract_cn: str = ""
    abstract_en: str = ""
    keywords: List[str] = field(default_factory=list)
    keywords_cn: List[str] = field(default_factory=list)
    keywords_en: List[str] = field(default_factory=list)
    body_items: List[Dict[str, Any]] = field(default_factory=list)
    sections: List[SectionBlock] = field(default_factory=list)
    tables: List[TableItem] = field(default_factory=list)
    references: List[str] = field(default_factory=list)
    page_count: int = 0
    image_count: int = 0


TRANSLATABLE_BODY_TYPES = {"heading", "paragraph", "figure"}

PAGE_TOP_MARGIN_CM = 2.5
PAGE_BOTTOM_MARGIN_CM = 2.2
PAGE_LEFT_MARGIN_CM = 2.3
PAGE_RIGHT_MARGIN_CM = 2.3
BODY_FONT_SIZE_PT = 10.5
BODY_LINE_SPACING = 1.22
BODY_FIRST_LINE_INDENT_CM = 0.74
HEADING_LEVEL_SIZES = {1: 12.0, 2: 11.0, 3: 10.5}


def build_requests_session() -> requests.Session:
    session = requests.Session()
    retry = Retry(
        total=3,
        connect=3,
        read=3,
        backoff_factor=1.0,
        status_forcelist=[429, 500, 502, 503, 504],
        allowed_methods=frozenset(["GET"]),
        raise_on_status=False,
    )
    adapter = HTTPAdapter(max_retries=retry)
    session.mount("http://", adapter)
    session.mount("https://", adapter)
    session.headers.update(
        {
            "User-Agent": "Flask-arXiv-Word-App/2.0 (structured-layout; contact: local-app)"
        }
    )
    return session


http_session = build_requests_session()


def safe_remove(path: str) -> None:
    try:
        if path and os.path.exists(path):
            os.remove(path)
    except Exception:
        pass


def safe_rmtree(path: str) -> None:
    try:
        if path and os.path.exists(path):
            shutil.rmtree(path, ignore_errors=True)
    except Exception:
        pass


def cleanup_generated_file(file_id: str) -> None:
    with GENERATED_FILES_LOCK:
        info = GENERATED_FILES.pop(file_id, None)

    if not info:
        return

    safe_rmtree(info.get("work_dir"))


def cleanup_expired_files() -> None:
    now = time.time()
    expired_ids = []

    with GENERATED_FILES_LOCK:
        for file_id, info in GENERATED_FILES.items():
            if now - info.get("created_at", 0) > FILE_TTL_SECONDS:
                expired_ids.append(file_id)

    for file_id in expired_ids:
        cleanup_generated_file(file_id)


@app.before_request
def before_request_cleanup() -> None:
    cleanup_expired_files()


def clean_whitespace(text: str) -> str:
    if not text:
        return ""
    text = html.unescape(text)
    text = text.replace("\r", "\n")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.split("\n")]
    text = "\n".join(lines)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def normalize_inline_text(text: str) -> str:
    text = html.unescape(text or "")
    text = text.replace("\xa0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def normalize_compare_text(text: str) -> str:
    text = clean_whitespace(text).lower()
    text = re.sub(r"\s+", "", text)
    text = re.sub(r"[^\w\u4e00-\u9fff]+", "", text)
    return text


def sanitize_filename(name: str, default: str = "paper") -> str:
    if not name:
        return default
    name = name.strip()
    name = re.sub(r'[\\/:*?"<>|\r\n]+', "_", name)
    name = re.sub(r"\s+", " ", name)
    name = name[:80].strip(" ._")
    return name or default


def split_keywords(text: str) -> List[str]:
    if not text:
        return []
    parts = [part.strip(" ;,.") for part in re.split(r"[;；]\s*|,\s*(?=[A-Z0-9])", text) if part.strip(" ;,.")]
    if len(parts) <= 1:
        parts = [part.strip(" ;,.") for part in re.split(r"[;,；，]\s*", text) if part.strip(" ;,.")]
    return parts


def contains_chinese(text: str) -> bool:
    return bool(re.search(r"[\u4e00-\u9fff]", text or ""))


def normalize_paragraph_text(text: str) -> str:
    text = html.unescape(text or "")
    text = text.replace("\r", "\n")
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    text = re.sub(r"(?<!\n)\n(?!\n)", " ", text)
    text = clean_whitespace(text)
    text = re.sub(r"\s+([,.;:!?])", r"\1", text)
    return text


def dedupe_keep_order(items: Iterable[str]) -> List[str]:
    seen = set()
    result = []
    for item in items:
        cleaned = clean_whitespace(item)
        key = normalize_compare_text(cleaned)
        if not cleaned or not key or key in seen:
            continue
        seen.add(key)
        result.append(cleaned)
    return result


def ollama_generate(prompt: str, model: str = OLLAMA_MODEL) -> str:
    try:
        response = http_session.post(
            OLLAMA_API_URL,
            json={
                "model": model,
                "prompt": prompt,
                "stream": False,
                "options": {
                    "temperature": 0.1,
                },
            },
            timeout=(15, 300),
        )
        response.raise_for_status()
    except requests.RequestException as e:
        raise RuntimeError(f"Ollama 翻译服务调用失败：{e}") from e

    data = response.json()
    translated = clean_whitespace(data.get("response", "") or "")
    if not translated:
        raise RuntimeError("Ollama 未返回可用翻译结果。")
    return translated


def build_translation_prompt(text: str) -> str:
    return (
        "你是一名专业学术翻译。请将下面的英文论文内容准确翻译成中文。\n"
        "要求：\n"
        "1. 仅输出中文翻译结果，不要解释。\n"
        "2. 保留原有的数学符号、公式、编号、引用编号和专有名词缩写。\n"
        "3. 保持学术正式语气。\n"
        "4. 如果原文已经是中文，则直接原样输出。\n\n"
        f"原文：\n{text}"
    )


def build_batch_translation_prompt(texts: List[str]) -> str:
    segments = TRANSLATION_SEPARATOR.join(texts)
    return (
        "你是一名专业学术翻译。请将以下多段英文论文内容逐段翻译成中文。\n"
        "必须严格遵守：\n"
        "1. 输出段落数量必须与输入一致。\n"
        f"2. 使用完全相同的分隔符 `{TRANSLATION_SEPARATOR}` 分隔每一段翻译结果。\n"
        "3. 不要添加任何解释、标题、前缀或后缀。\n"
        "4. 保留公式、变量名、引用编号、术语缩写。\n"
        "5. 如果某段已经是中文，则直接原样输出。\n\n"
        f"待翻译内容：\n{segments}"
    )


def clean_translation_output(text: str) -> str:
    cleaned = clean_whitespace(text)
    cleaned = cleaned.replace(TRANSLATION_SEPARATOR, "\n")
    cleaned = cleaned.replace("<<<SEG>>>", "")
    return clean_whitespace(cleaned)


def translate_segments(texts: List[str], model: str = OLLAMA_MODEL) -> List[str]:
    if not texts:
        return []

    cleaned_texts = [clean_whitespace(text) for text in texts]
    result = [""] * len(cleaned_texts)
    pending_indices = [index for index, text in enumerate(cleaned_texts) if text and not contains_chinese(text)]

    for index, text in enumerate(cleaned_texts):
        if not text:
            result[index] = ""
        elif contains_chinese(text):
            result[index] = text

    batch: List[Tuple[int, str]] = []
    current_len = 0
    batched_groups: List[List[Tuple[int, str]]] = []
    for index in pending_indices:
        text = cleaned_texts[index]
        projected = current_len + len(text)
        if batch and projected > 2400:
            batched_groups.append(batch)
            batch = []
            current_len = 0
        batch.append((index, text))
        current_len += len(text)
    if batch:
        batched_groups.append(batch)

    for group in batched_groups:
        texts_in_group = [text for _, text in group]
        translated_batch = None
        if len(texts_in_group) == 1:
            translated_batch = [clean_translation_output(ollama_generate(build_translation_prompt(texts_in_group[0]), model=model))]
        else:
            try:
                translated_raw = ollama_generate(build_batch_translation_prompt(texts_in_group), model=model)
                parts = translated_raw.split(TRANSLATION_SEPARATOR)
                if len(parts) == len(texts_in_group):
                    translated_batch = [clean_translation_output(part) for part in parts]
            except Exception:
                translated_batch = None

        if translated_batch is None:
            translated_batch = [
                clean_translation_output(ollama_generate(build_translation_prompt(text), model=model))
                for text in texts_in_group
            ]

        for (index, _), translated in zip(group, translated_batch):
            result[index] = translated or cleaned_texts[index]

    return result


def get_arxiv_pdf_url(entry: Any) -> str:
    links = getattr(entry, "links", []) or []
    for link in links:
        href = getattr(link, "href", "") or ""
        title = getattr(link, "title", "") or ""
        link_type = getattr(link, "type", "") or ""
        if title.lower() == "pdf" or link_type == "application/pdf" or href.endswith(".pdf"):
            return href

    entry_id = getattr(entry, "id", "") or ""
    if "/abs/" in entry_id:
        arxiv_id = entry_id.rsplit("/abs/", 1)[-1]
        return f"https://arxiv.org/pdf/{arxiv_id}.pdf"

    raise ValueError("未能解析该论文的 PDF 链接。")


def search_arxiv(keyword: str, max_results: int = 50) -> List[Dict[str, Any]]:
    if not keyword.strip():
        raise ValueError("请输入检索关键词。")

    resp = http_session.get(
        "http://export.arxiv.org/api/query",
        params={"search_query": f"all:{keyword}", "start": 0, "max_results": max_results},
        timeout=(8, 25),
    )
    resp.raise_for_status()

    parsed = feedparser.parse(resp.text)
    results: List[Dict[str, Any]] = []

    for entry in parsed.entries:
        title = clean_whitespace(getattr(entry, "title", "") or "无标题")
        summary = clean_whitespace(getattr(entry, "summary", "") or "")
        published = getattr(entry, "published", "") or getattr(entry, "updated", "") or ""
        year = published[:4] if len(published) >= 4 else "未知"

        authors = []
        if getattr(entry, "authors", None):
            for author in entry.authors:
                name = clean_whitespace(getattr(author, "name", "") or "")
                if name:
                    authors.append(name)
        elif getattr(entry, "author", None):
            authors = [clean_whitespace(x) for x in str(entry.author).split(",") if clean_whitespace(x)]

        results.append(
            {
                "title": title,
                "authors": authors,
                "year": year,
                "summary": summary,
                "pdf_url": get_arxiv_pdf_url(entry),
            }
        )

    return results


def download_pdf(pdf_url: str, save_path: str) -> None:
    try:
        resp = http_session.get(pdf_url, stream=True, timeout=(10, 60))
        resp.raise_for_status()
    except requests.RequestException as e:
        raise RuntimeError(f"PDF 下载失败：{e}") from e

    with open(save_path, "wb") as f:
        for chunk in resp.iter_content(chunk_size=8192):
            if chunk:
                f.write(chunk)


def is_section_heading(text: str, font_size: float, bold: bool) -> bool:
    cleaned = clean_whitespace(text)
    if not cleaned or len(cleaned) > 120:
        return False

    upper = cleaned.upper()
    if REFERENCE_HEADING_RE.match(cleaned):
        return True
    if ABSTRACT_RE.match(cleaned):
        return True
    if re.match(r"^(?:[IVXLC]+\.?\s+)?(INTRODUCTION|RELATED WORK|BACKGROUND|METHODS?|METHODOLOGY|EXPERIMENTS?|RESULTS?|DISCUSSION|CONCLUSION|CONCLUSIONS|ACKNOWLEDGMENTS?)$", upper):
        return True
    if re.match(r"^(?:\d+(?:\.\d+)*)\s+[A-Z\u4e00-\u9fff]", cleaned):
        return True
    if re.match(r"^(?:[IVXLC]+[.)]?\s+)[A-Z]", cleaned):
        return True
    if re.match(r"^[A-Z][A-Z0-9 ,:/\-]{3,}$", cleaned) and len(cleaned.split()) <= 10 and (bold or font_size >= 11.0):
        return True
    if re.match(r"^\d+\s*[\u4e00-\u9fffA-Za-z].{0,40}$", cleaned) and (bold or font_size >= 11.0):
        return True
    return bool(bold and font_size >= 11.5 and len(cleaned.split()) <= 12 and not cleaned.endswith("."))


def is_reference_item_start(text: str) -> bool:
    cleaned = clean_whitespace(text)
    return bool(
        re.match(r"^\[\d+\]\s+", cleaned)
        or re.match(r"^\d+\.\s+", cleaned)
        or re.match(r"^\[[A-Za-z0-9]+\]\s+", cleaned)
    )


def extract_text_blocks(pdf_doc: fitz.Document) -> Tuple[List[TextBlock], Dict[int, float]]:
    blocks: List[TextBlock] = []
    page_heights: Dict[int, float] = {}

    for page_index in range(pdf_doc.page_count):
        page = pdf_doc.load_page(page_index)
        page_heights[page_index + 1] = float(page.rect.height)
        raw = page.get_text("dict")

        for block_number, block in enumerate(raw.get("blocks", [])):
            if block.get("type") != 0:
                continue

            block_lines = []
            font_sizes: List[float] = []
            bold_flags: List[bool] = []

            for line in block.get("lines", []):
                line_parts = []
                for span in line.get("spans", []):
                    span_text = normalize_inline_text(span.get("text", ""))
                    if not span_text:
                        continue
                    line_parts.append(span_text)

                    font_sizes.append(float(span.get("size", 0.0) or 0.0))
                    font_name = str(span.get("font", "") or "").lower()
                    flags = int(span.get("flags", 0) or 0)
                    bold_flags.append("bold" in font_name or "black" in font_name or bool(flags & 16))

                line_text = normalize_inline_text(" ".join(line_parts))
                if line_text:
                    block_lines.append(line_text)

            text = clean_whitespace("\n".join(block_lines))
            if not text:
                continue

            blocks.append(
                TextBlock(
                    id=f"p{page_index + 1}_b{block_number}",
                    page_number=page_index + 1,
                    block_number=block_number,
                    bbox=tuple(block.get("bbox", (0, 0, 0, 0))),
                    text=text,
                    font_size=max(font_sizes) if font_sizes else 0.0,
                    bold=bold_flags.count(True) >= max(1, len(bold_flags) // 2),
                    line_count=max(1, len(block_lines)),
                )
            )

    blocks.sort(key=lambda b: (b.page_number, b.y0, b.x0))
    return blocks, page_heights


def detect_repeated_margin_noise(blocks: List[TextBlock], page_heights: Dict[int, float]) -> Set[str]:
    occurrences: Dict[str, Set[int]] = defaultdict(set)
    for block in blocks:
        page_height = page_heights.get(block.page_number, 0.0)
        in_margin = block.y0 < page_height * 0.12 or block.y1 > page_height * 0.88
        if not in_margin:
            continue

        normalized = normalize_compare_text(block.text)
        if not normalized or len(normalized) < 4 or len(normalized) > 120:
            continue
        occurrences[normalized].add(block.page_number)

    return {text for text, pages in occurrences.items() if len(pages) >= 2}


def is_noise_block(block: TextBlock, repeated_margin_noise: Set[str], paper_title: str, page_heights: Dict[int, float]) -> bool:
    text = clean_whitespace(block.text)
    normalized = normalize_compare_text(text)
    page_height = page_heights.get(block.page_number, 0.0)
    in_margin = block.y0 < page_height * 0.12 or block.y1 > page_height * 0.88
    lower = text.lower()

    if not normalized:
        return True
    if normalized in repeated_margin_noise:
        return True
    if ARXIV_NOISE_RE.search(lower):
        return True
    if in_margin and normalize_compare_text(paper_title) == normalized:
        return True
    if in_margin and re.match(r"^(page\s*)?\d+$", lower):
        return True
    if in_margin and re.match(r"^\d+\s*/\s*\d+$", lower):
        return True
    if in_margin and lower.startswith("preprint"):
        return True
    if in_margin and re.match(r"^\(?\d+\)?$", text):
        return True
    if in_margin and len(normalized) <= 2:
        return True
    return False


def extract_images(pdf_doc: fitz.Document, temp_dir: str) -> List[Dict[str, Any]]:
    image_records: List[Dict[str, Any]] = []

    for page_index in range(pdf_doc.page_count):
        page = pdf_doc.load_page(page_index)
        raw = page.get_text("dict")
        image_boxes = [tuple(block.get("bbox", (0, 0, 0, 0))) for block in raw.get("blocks", []) if block.get("type") == 1]
        xrefs = []
        seen = set()

        for image in page.get_images(full=True):
            xref = image[0]
            if xref in seen:
                continue
            seen.add(xref)
            xrefs.append(xref)

        for index, xref in enumerate(xrefs, start=1):
            try:
                pix = fitz.Pixmap(pdf_doc, xref)
                if pix.n - pix.alpha > 3:
                    pix = fitz.Pixmap(fitz.csRGB, pix)

                image_path = os.path.join(temp_dir, f"page_{page_index + 1}_img_{index}.png")
                pix.save(image_path)
                pix = None
            except Exception:
                continue

            bbox = image_boxes[index - 1] if index - 1 < len(image_boxes) else (0.0, 0.0, page.rect.width * 0.7, page.rect.height * 0.3)
            width_ratio = max(0.1, min(1.0, (bbox[2] - bbox[0]) / max(1.0, page.rect.width)))
            image_records.append(
                {
                    "page_number": page_index + 1,
                    "bbox": bbox,
                    "path": image_path,
                    "width_ratio": width_ratio,
                }
            )

    return image_records


def associate_figures(blocks: List[TextBlock], image_records: List[Dict[str, Any]]) -> Tuple[List[FigureItem], Set[str]]:
    blocks_by_page: Dict[int, List[TextBlock]] = defaultdict(list)
    for block in blocks:
        blocks_by_page[block.page_number].append(block)

    figure_items: List[FigureItem] = []
    caption_block_ids: Set[str] = set()
    fallback_index = 1

    for image in image_records:
        page_blocks = blocks_by_page.get(image["page_number"], [])
        bbox = image["bbox"]
        below_candidates = []
        above_candidates = []

        for block in page_blocks:
            if FIGURE_CAPTION_RE.match(block.text):
                if 0 <= block.y0 - bbox[3] <= 120:
                    below_candidates.append(block)
                elif 0 <= bbox[1] - block.y1 <= 60:
                    above_candidates.append(block)

        caption_block = None
        if below_candidates:
            caption_block = min(below_candidates, key=lambda b: b.y0 - bbox[3])
        elif above_candidates:
            caption_block = min(above_candidates, key=lambda b: bbox[1] - b.y1)

        if caption_block is not None:
            caption = normalize_paragraph_text(caption_block.text)
            caption_block_ids.add(caption_block.id)
            order_y = min(float(bbox[1]), caption_block.y0)
        else:
            caption = f"Figure {fallback_index}."
            order_y = float(bbox[1])

        figure_items.append(
            FigureItem(
                page_number=image["page_number"],
                order_y=order_y,
                image_path=image["path"],
                caption=caption,
                width_ratio=float(image["width_ratio"]),
                is_wide=float(image["width_ratio"]) >= 0.72,
            )
        )
        fallback_index += 1

    figure_items.sort(key=lambda item: (item.page_number, item.order_y))
    return figure_items, caption_block_ids


def infer_title(text_blocks: List[TextBlock], paper: Dict[str, Any]) -> str:
    provided_title = clean_whitespace(paper.get("title", "") or "")
    if provided_title:
        return provided_title

    first_page_blocks = [block for block in text_blocks if block.page_number == 1 and block.y0 < 220]
    if not first_page_blocks:
        return "Untitled Paper"

    title_block = max(first_page_blocks, key=lambda block: (block.font_size, -(block.y0)))
    return normalize_paragraph_text(title_block.text) or "Untitled Paper"


def find_heading_index(blocks: List[TextBlock], predicate) -> Optional[int]:
    for index, block in enumerate(blocks):
        if predicate(block):
            return index
    return None


def infer_authors_and_affiliations(
    first_page_blocks: List[TextBlock],
    title: str,
    provided_authors: List[str],
    abstract_index: Optional[int],
) -> Tuple[List[str], List[str], Set[str]]:
    consumed_ids: Set[str] = set()
    authors = dedupe_keep_order(provided_authors or [])
    affiliations: List[str] = []

    title_norm = normalize_compare_text(title)
    title_bottom = 0.0
    title_ids = set()
    for block in first_page_blocks:
        if title_norm and title_norm in normalize_compare_text(block.text):
            title_bottom = max(title_bottom, block.y1)
            title_ids.add(block.id)

    cutoff_y = first_page_blocks[-1].y1 if first_page_blocks else 0.0
    if abstract_index is not None and 0 <= abstract_index < len(first_page_blocks):
        cutoff_y = first_page_blocks[abstract_index].y0

    candidate_blocks = [
        block
        for block in first_page_blocks
        if block.id not in title_ids and block.y0 >= max(0.0, title_bottom - 4.0) and block.y1 <= cutoff_y
    ]

    author_block_ids = set()
    if not authors:
        for block in candidate_blocks:
            text = normalize_paragraph_text(block.text)
            lower = text.lower()
            if "@" in text:
                continue
            if any(token in lower for token in ["university", "institute", "college", "department", "laboratory", "school", "academy", "center", "centre"]):
                continue
            if len(text.split()) <= 20 and ("," in text or " and " in lower or len(text.split()) <= 8):
                possible_authors = re.split(r",| and ", text)
                possible_authors = [clean_whitespace(item) for item in possible_authors if clean_whitespace(item)]
                if possible_authors:
                    authors = possible_authors
                    author_block_ids.add(block.id)
                    break

    if authors:
        normalized_authors = {normalize_compare_text(name) for name in authors}
        for block in candidate_blocks:
            block_norm = normalize_compare_text(block.text)
            if any(name and name in block_norm for name in normalized_authors):
                author_block_ids.add(block.id)

    for block in candidate_blocks:
        if block.id in author_block_ids:
            continue
        text = normalize_paragraph_text(block.text)
        lower = text.lower()
        if any(
            token in lower
            for token in ["university", "institute", "college", "department", "laboratory", "school", "academy", "center", "centre", "@", "hospital"]
        ):
            affiliations.append(text)
            consumed_ids.add(block.id)
        elif re.search(r"\b[A-Z][a-z]+ University\b", text):
            affiliations.append(text)
            consumed_ids.add(block.id)

    consumed_ids.update(author_block_ids)
    affiliations = dedupe_keep_order(affiliations)
    return authors, affiliations, consumed_ids


def parse_abstract_and_keywords(
    blocks: List[TextBlock],
    abstract_index: Optional[int],
    reference_index: Optional[int],
    fallback_summary: str,
) -> Tuple[str, List[str], int, Set[str]]:
    if abstract_index is None:
        return clean_whitespace(fallback_summary or ""), [], 0, set()

    abstract_parts: List[str] = []
    keywords: List[str] = []
    consumed_ids: Set[str] = set()
    last_index = abstract_index

    first_match = ABSTRACT_RE.match(blocks[abstract_index].text)
    consumed_ids.add(blocks[abstract_index].id)
    if first_match and first_match.group(2).strip():
        abstract_parts.append(normalize_paragraph_text(first_match.group(2)))

    for index in range(abstract_index + 1, len(blocks)):
        block = blocks[index]
        if reference_index is not None and index >= reference_index:
            break
        if index > abstract_index + 1 and is_section_heading(block.text, block.font_size, block.bold):
            break

        keyword_match = KEYWORDS_RE.match(block.text)
        if keyword_match:
            keywords.extend(split_keywords(keyword_match.group(2)))
            consumed_ids.add(block.id)
            last_index = index
            break

        abstract_parts.append(normalize_paragraph_text(block.text))
        consumed_ids.add(block.id)
        last_index = index

    if not abstract_parts:
        abstract_parts = [clean_whitespace(fallback_summary or "")]

    abstract_text = " ".join(part for part in abstract_parts if part).strip()
    abstract_text = clean_whitespace(abstract_text)

    if not keywords:
        keyword_match = KEYWORDS_RE.search(abstract_text)
        if keyword_match:
            abstract_text = clean_whitespace(abstract_text[: keyword_match.start()])
            keywords = split_keywords(keyword_match.group(2))

    return abstract_text, dedupe_keep_order(keywords), last_index, consumed_ids


def build_references(blocks: List[TextBlock]) -> List[str]:
    refs: List[str] = []
    current = ""

    for block in blocks:
        paragraphs = [normalize_paragraph_text(p) for p in block.text.split("\n") if normalize_paragraph_text(p)]
        for paragraph in paragraphs:
            if is_reference_item_start(paragraph):
                if current:
                    refs.append(current.strip())
                current = paragraph
            else:
                if current:
                    current = f"{current} {paragraph}".strip()
                else:
                    current = paragraph

    if current:
        refs.append(current.strip())

    if not refs:
        refs = [normalize_paragraph_text(block.text) for block in blocks if normalize_paragraph_text(block.text)]

    return refs


def detect_funding_or_notes(first_page_blocks: List[TextBlock]) -> Tuple[List[str], List[str], Set[str]]:
    funding: List[str] = []
    notes: List[str] = []
    consumed: Set[str] = set()

    for block in first_page_blocks:
        text = normalize_paragraph_text(block.text)
        lower = text.lower()
        if any(token in lower for token in ["fund", "supported by", "grant", "基金", "资助", "项目"]):
            funding.append(text)
            consumed.add(block.id)
        elif any(token in lower for token in ["corresponding author", "通讯作者", "author biography", "作者简介", "@"]):
            notes.append(text)
            consumed.add(block.id)

    return dedupe_keep_order(funding), dedupe_keep_order(notes), consumed


def infer_heading_level(text: str) -> int:
    cleaned = clean_whitespace(text)
    if REFERENCE_HEADING_RE.match(cleaned):
        return 1
    if re.match(r"^(?:[IVXLC]+[.)]?\s+)", cleaned):
        return 1
    if re.match(r"^\d+\s+", cleaned):
        return 1
    if re.match(r"^\d+\.\d+\s+", cleaned):
        return 2
    if re.match(r"^\d+\.\d+\.\d+\s+", cleaned):
        return 3
    if re.match(r"^[A-Z][A-Z0-9 ,:/\-]{3,}$", cleaned):
        return 1
    return 2


def looks_like_table_block(text: str) -> bool:
    lines = [clean_whitespace(line) for line in text.split("\n") if clean_whitespace(line)]
    if len(lines) < 2:
        return False
    delimiter_like = 0
    for line in lines:
        if "\t" in line or "|" in line:
            delimiter_like += 1
        elif len(re.findall(r"\s{2,}", line)) >= 1:
            delimiter_like += 1
    return delimiter_like >= max(2, len(lines) // 2)


def split_table_row(line: str) -> List[str]:
    if "|" in line:
        parts = [part.strip() for part in line.split("|") if part.strip()]
    elif "\t" in line:
        parts = [part.strip() for part in line.split("\t") if part.strip()]
    else:
        parts = [part.strip() for part in re.split(r"\s{2,}", line) if part.strip()]
    return parts


def detect_table_caption(text: str) -> Optional[str]:
    cleaned = clean_whitespace(text)
    if re.match(r"^(table|表)\s*\d+[:：.\-]?\s*.+$", cleaned, re.I):
        return cleaned
    return None


def build_sections_from_body_items(body_items: List[Dict[str, Any]]) -> List[SectionBlock]:
    sections: List[SectionBlock] = []
    current_section = SectionBlock(heading="正文", level=1, items=[])

    for item in body_items:
        item_type = item.get("type")
        if item_type == "heading":
            if current_section.items or current_section.heading != "正文":
                sections.append(current_section)
            current_section = SectionBlock(
                heading=item["text"],
                level=infer_heading_level(item["text"]),
                items=[],
            )
            continue

        if item_type == "paragraph":
            current_section.items.append(SectionItem(type="paragraph", text=item["text"]))
        elif item_type == "figure":
            current_section.items.append(SectionItem(type="figure", figure=item["figure"]))
        elif item_type == "table":
            current_section.items.append(SectionItem(type="table", table=item["table"]))

    if current_section.items or current_section.heading:
        sections.append(current_section)

    return sections


def populate_bilingual_metadata(article: ArticleStructure) -> ArticleStructure:
    if article.title_cn or article.title_en or article.abstract_cn or article.abstract_en:
        return article

    if contains_chinese(article.title):
        article.title_cn = article.title
    else:
        article.title_en = article.title

    if contains_chinese(article.abstract):
        article.abstract_cn = article.abstract
        article.keywords_cn = list(article.keywords)
    else:
        article.abstract_en = article.abstract
        article.keywords_en = list(article.keywords)

    return article


def translate_article_to_chinese(article: ArticleStructure, model: str = OLLAMA_MODEL) -> ArticleStructure:
    translated = ArticleStructure(
        title=article.title,
        title_cn=article.title_cn,
        title_en=article.title_en,
        authors=list(article.authors),
        affiliations=list(article.affiliations),
        funding=list(article.funding),
        author_notes=list(article.author_notes),
        abstract=article.abstract,
        abstract_cn=article.abstract_cn,
        abstract_en=article.abstract_en,
        keywords=list(article.keywords),
        keywords_cn=list(article.keywords_cn),
        keywords_en=list(article.keywords_en),
        body_items=[],
        sections=[],
        tables=list(article.tables),
        references=list(article.references),
        page_count=article.page_count,
        image_count=article.image_count,
    )

    original = populate_bilingual_metadata(article)
    translated.title_en = original.title_en or original.title
    translated.abstract_en = original.abstract_en or original.abstract
    translated.keywords_en = list(original.keywords_en or original.keywords)

    translated.title = translate_segments([article.title], model=model)[0] if article.title else article.title
    translated.title_cn = translated.title
    if article.affiliations:
        translated.affiliations = translate_segments(article.affiliations, model=model)
    if article.funding:
        translated.funding = translate_segments(article.funding, model=model)
    if article.author_notes:
        translated.author_notes = translate_segments(article.author_notes, model=model)
    if article.abstract:
        translated.abstract = translate_segments([article.abstract], model=model)[0]
        translated.abstract_cn = translated.abstract
    if article.keywords:
        translated.keywords = translate_segments(article.keywords, model=model)
        translated.keywords_cn = list(translated.keywords)

    translatable_indices: List[int] = []
    translatable_texts: List[str] = []
    translated.body_items = [dict(item) for item in article.body_items]

    for index, item in enumerate(article.body_items):
        item_type = item.get("type")
        if item_type == "paragraph":
            translatable_indices.append(index)
            translatable_texts.append(item["text"])
        elif item_type == "heading":
            translatable_indices.append(index)
            translatable_texts.append(item["text"])
        elif item_type == "figure":
            figure: FigureItem = item["figure"]
            translatable_indices.append(index)
            translatable_texts.append(figure.caption)

    translated_texts = translate_segments(translatable_texts, model=model)

    for item_index, translated_text in zip(translatable_indices, translated_texts):
        item = translated.body_items[item_index]
        if item["type"] in {"paragraph", "heading"}:
            item["text"] = translated_text
        elif item["type"] == "figure":
            figure = item["figure"]
            item["figure"] = FigureItem(
                page_number=figure.page_number,
                order_y=figure.order_y,
                image_path=figure.image_path,
                caption=translated_text,
                width_ratio=figure.width_ratio,
                is_wide=figure.is_wide,
            )

    translated.sections = build_sections_from_body_items(translated.body_items)
    return translated


def parse_pdf_to_article(pdf_path: str, temp_dir: str, paper: Dict[str, Any]) -> ArticleStructure:
    try:
        pdf_doc = fitz.open(pdf_path)
    except Exception as e:
        raise RuntimeError(f"PDF 解析失败：{e}") from e

    try:
        text_blocks, page_heights = extract_text_blocks(pdf_doc)
        repeated_margin_noise = detect_repeated_margin_noise(text_blocks, page_heights)
        filtered_blocks = [
            block
            for block in text_blocks
            if not is_noise_block(block, repeated_margin_noise, paper.get("title", ""), page_heights)
        ]
        figure_items, caption_block_ids = associate_figures(filtered_blocks, extract_images(pdf_doc, temp_dir))

        title = infer_title(filtered_blocks, paper)
        blocks_by_page: Dict[int, List[TextBlock]] = defaultdict(list)
        for block in filtered_blocks:
            blocks_by_page[block.page_number].append(block)

        first_page_blocks = blocks_by_page.get(1, [])
        abstract_index_first_page = find_heading_index(first_page_blocks, lambda block: bool(ABSTRACT_RE.match(block.text)))
        authors, affiliations, front_ids = infer_authors_and_affiliations(
            first_page_blocks=first_page_blocks,
            title=title,
            provided_authors=paper.get("authors", []) or [],
            abstract_index=abstract_index_first_page,
        )
        funding, author_notes, note_ids = detect_funding_or_notes(first_page_blocks)

        abstract_index = find_heading_index(filtered_blocks, lambda block: bool(ABSTRACT_RE.match(block.text)))
        reference_index = find_heading_index(filtered_blocks, lambda block: REFERENCE_HEADING_RE.match(block.text))

        abstract_text, keywords, abstract_end_index, abstract_ids = parse_abstract_and_keywords(
            blocks=filtered_blocks,
            abstract_index=abstract_index,
            reference_index=reference_index,
            fallback_summary=paper.get("summary", ""),
        )

        consumed_ids = set(front_ids)
        consumed_ids.update(note_ids)
        consumed_ids.update(abstract_ids)
        consumed_ids.update(caption_block_ids)

        body_start_index = abstract_end_index + 1 if abstract_index is not None else 0
        while body_start_index < len(filtered_blocks) and filtered_blocks[body_start_index].id in consumed_ids:
            body_start_index += 1

        ref_start = reference_index if reference_index is not None else len(filtered_blocks)
        id_to_index = {block.id: index for index, block in enumerate(filtered_blocks)}

        body_items: List[Dict[str, Any]] = []
        tables: List[TableItem] = []
        page_figures: Dict[int, List[FigureItem]] = defaultdict(list)
        for figure in figure_items:
            page_figures[figure.page_number].append(figure)

        for page_number in sorted({block.page_number for block in filtered_blocks} | set(page_figures.keys())):
            page_entries: List[Tuple[float, int, Dict[str, Any]]] = []

            for block in filtered_blocks:
                if block.page_number != page_number:
                    continue
                if block.id in consumed_ids:
                    continue

                block_index = id_to_index[block.id]
                if block_index < body_start_index or block_index >= ref_start:
                    continue

                text = normalize_paragraph_text(block.text)
                if not text:
                    continue

                item_type = "heading" if is_section_heading(text, block.font_size, block.bold) else "paragraph"
                if item_type == "paragraph" and looks_like_table_block(block.text):
                    rows = [split_table_row(line) for line in block.text.split("\n") if clean_whitespace(line)]
                    rows = [row for row in rows if row]
                    if len(rows) >= 2 and max(len(row) for row in rows) >= 2:
                        table_caption = detect_table_caption(rows[0][0]) if rows and rows[0] else None
                        if table_caption:
                            rows = rows[1:]
                        table_item = TableItem(
                            caption=table_caption or f"表{len(tables) + 1}",
                            rows=rows,
                            page_number=page_number,
                            order_y=block.y0,
                        )
                        tables.append(table_item)
                        page_entries.append((block.y0, 0, {"type": "table", "table": table_item}))
                        continue

                page_entries.append((block.y0, 0, {"type": item_type, "text": text, "page_number": page_number}))

            for figure in page_figures.get(page_number, []):
                page_entries.append((figure.order_y, 1, {"type": "figure", "figure": figure}))

            page_entries.sort(key=lambda entry: (entry[0], entry[1]))
            body_items.extend([item for _, _, item in page_entries])

        references = build_references(filtered_blocks[ref_start + 1 :]) if ref_start < len(filtered_blocks) else []

        article = ArticleStructure(
            title=title,
            authors=authors,
            affiliations=affiliations,
            funding=funding,
            author_notes=author_notes,
            abstract=abstract_text or clean_whitespace(paper.get("summary", "") or ""),
            keywords=keywords,
            body_items=body_items,
            tables=tables,
            references=references,
            page_count=pdf_doc.page_count,
            image_count=len(figure_items),
        )
        article.sections = build_sections_from_body_items(article.body_items)
        return populate_bilingual_metadata(article)
    finally:
        pdf_doc.close()


def set_run_font(
    run,
    latin_font: str,
    east_asia_font: str,
    font_size: float,
    bold: bool = False,
    italic: bool = False,
) -> None:
    run.font.name = latin_font
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)

    r_fonts.set(qn("w:ascii"), latin_font)
    r_fonts.set(qn("w:hAnsi"), latin_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)
    r_fonts.set(qn("w:cs"), latin_font)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, BODY_FONTS[0], BODY_FONTS[1], 9)


def set_section_columns(section, num: int, space_cm: float = 0.8) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    cols_el = cols[0] if cols else OxmlElement("w:cols")
    cols_el.set(qn("w:num"), str(max(1, num)))
    cols_el.set(qn("w:space"), str(int((space_cm / 2.54) * 1440)))
    if not cols:
        sect_pr.append(cols_el)


def set_document_margins(section) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(PAGE_TOP_MARGIN_CM)
    section.bottom_margin = Cm(PAGE_BOTTOM_MARGIN_CM)
    section.left_margin = Cm(PAGE_LEFT_MARGIN_CM)
    section.right_margin = Cm(PAGE_RIGHT_MARGIN_CM)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)


def configure_section(section, columns: int = 1) -> None:
    set_document_margins(section)
    set_section_columns(section, columns)

    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.text = ""
    add_page_number(paragraph)


def create_academic_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = BODY_FONTS[0]
    normal.font.size = Pt(BODY_FONT_SIZE_PT)
    r_pr = normal._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), BODY_FONTS[0])
    r_fonts.set(qn("w:hAnsi"), BODY_FONTS[0])
    r_fonts.set(qn("w:eastAsia"), BODY_FONTS[1])

    for section in document.sections:
        configure_section(section, columns=1)


def add_english_title(document: Document, title_en: str) -> None:
    if not clean_whitespace(title_en):
        return
    add_paragraph(
        document,
        title_en,
        latin_font="Times New Roman",
        east_asia_font=BODY_FONTS[1],
        font_size=13.5,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_cm=0.0,
        space_before_pt=0.0,
        space_after_pt=8.0,
        line_spacing=1.0,
    )


def add_small_note_block(document: Document, title: str, items: List[str]) -> None:
    if not items:
        return
    add_paragraph(
        document,
        title,
        latin_font=HEADING_FONTS[0],
        east_asia_font=HEADING_FONTS[1],
        font_size=9.0,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent_cm=0.0,
        space_before_pt=4.0,
        space_after_pt=1.0,
        line_spacing=1.0,
    )
    for item in items:
        add_paragraph(
            document,
            item,
            font_size=9.0,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent_cm=0.0,
            space_before_pt=0.0,
            space_after_pt=1.0,
            line_spacing=1.0,
        )


def add_keywords_block(document: Document, label: str, keywords: List[str], chinese_mode: bool) -> None:
    if not keywords:
        return
    add_label_block(document, label, "; ".join(keywords), chinese_mode=chinese_mode)


def add_abstract_block(document: Document, label: str, content: str, chinese_mode: bool) -> None:
    if not clean_whitespace(content):
        return
    add_label_block(document, label, content, chinese_mode=chinese_mode)


def add_title_block(document: Document, article: ArticleStructure) -> None:
    title_cn = clean_whitespace(article.title_cn)
    title_en = clean_whitespace(article.title_en)
    primary_title = title_cn or clean_whitespace(article.title)

    add_paragraph(
        document,
        primary_title,
        latin_font=TITLE_FONTS[0],
        east_asia_font=TITLE_FONTS[1],
        font_size=16,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_cm=0.0,
        space_before_pt=6.0,
        space_after_pt=8.0,
        line_spacing=1.0,
    )
    add_english_title(document, title_en if title_en and title_en != primary_title else "")

    if article.authors:
        add_paragraph(
            document,
            " / ".join(article.authors) if contains_chinese(primary_title) else ", ".join(article.authors),
            font_size=11.0,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent_cm=0.0,
            space_before_pt=0.0,
            space_after_pt=4.0,
            line_spacing=1.0,
        )

    for affiliation in article.affiliations:
        add_paragraph(
            document,
            affiliation,
            font_size=10.0,
            italic=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent_cm=0.0,
            space_before_pt=0.0,
            space_after_pt=2.0,
            line_spacing=1.0,
        )


def add_paragraph(
    container,
    text: str = "",
    *,
    latin_font: str = BODY_FONTS[0],
    east_asia_font: str = BODY_FONTS[1],
    font_size: float = 10.5,
    bold: bool = False,
    italic: bool = False,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent_cm: float = 0.74,
    left_indent_cm: float = 0.0,
    hanging_indent_cm: float = 0.0,
    space_before_pt: float = 0.0,
    space_after_pt: float = 4.0,
    line_spacing: float = 1.15,
):
    paragraph = container.add_paragraph()
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(space_before_pt)
    fmt.space_after = Pt(space_after_pt)
    fmt.line_spacing = line_spacing
    fmt.left_indent = Cm(left_indent_cm)
    if hanging_indent_cm:
        fmt.first_line_indent = Cm(-abs(hanging_indent_cm))
    else:
        fmt.first_line_indent = Cm(first_line_indent_cm)

    run = paragraph.add_run(text)
    set_run_font(run, latin_font, east_asia_font, font_size, bold=bold, italic=italic)
    return paragraph


def add_label_block(container, label: str, content: str, *, chinese_mode: bool) -> None:
    paragraph = container.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(2)
    fmt.space_after = Pt(6)
    fmt.line_spacing = 1.15
    fmt.first_line_indent = Cm(0)

    label_run = paragraph.add_run(label)
    set_run_font(label_run, BODY_FONTS[0], BODY_FONTS[1], 10.5, bold=True)
    content_run = paragraph.add_run(f" {content}" if not chinese_mode else content)
    set_run_font(content_run, BODY_FONTS[0], BODY_FONTS[1], 10.5)


def add_heading_paragraph(container, text: str, level: int = 1) -> None:
    paragraph = add_paragraph(
        container,
        text,
        latin_font=HEADING_FONTS[0],
        east_asia_font=HEADING_FONTS[1],
        font_size=HEADING_LEVEL_SIZES.get(level, 10.5),
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent_cm=0.0,
        space_before_pt=8.0 if level == 1 else 6.0,
        space_after_pt=4.0,
        line_spacing=1.0,
    )
    paragraph.paragraph_format.keep_with_next = True


def add_figure_block(document: Document, figure: FigureItem, *, single_column: bool) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run()
    run.add_picture(figure.image_path, width=Cm(16.0 if single_column else 7.2))

    add_paragraph(
        document,
        figure.caption,
        latin_font=CAPTION_FONTS[0],
        east_asia_font=CAPTION_FONTS[1],
        font_size=9.0,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_cm=0.0,
        space_before_pt=0.0,
        space_after_pt=6.0,
        line_spacing=1.0,
    )


def add_front_matter(document: Document, article: ArticleStructure) -> None:
    add_title_block(document, article)
    document.add_paragraph()
    add_abstract_block(document, "摘要：", article.abstract_cn, True)
    add_keywords_block(document, "关键词：", article.keywords_cn, True)
    add_abstract_block(document, "Abstract:", article.abstract_en, False)
    add_keywords_block(document, "Key words:", article.keywords_en, False)
    add_small_note_block(document, "基金项目", article.funding)
    add_small_note_block(document, "作者说明", article.author_notes)


def add_table_with_caption(document: Document, table_item: TableItem, single_column: bool) -> None:
    if table_item.caption:
        add_paragraph(
            document,
            table_item.caption,
            latin_font=CAPTION_FONTS[0],
            east_asia_font=CAPTION_FONTS[1],
            font_size=9.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent_cm=0.0,
            space_before_pt=4.0,
            space_after_pt=4.0,
            line_spacing=1.0,
        )

    if not table_item.rows:
        return

    column_count = max(len(row) for row in table_item.rows)
    table = document.add_table(rows=0, cols=column_count)
    table.style = "Table Grid"
    for row in table_item.rows:
        cells = table.add_row().cells
        for index in range(column_count):
            cells[index].text = row[index] if index < len(row) else ""
            for paragraph in cells[index].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(run, BODY_FONTS[0], BODY_FONTS[1], 9.5)

    add_paragraph(
        document,
        "",
        first_line_indent_cm=0.0,
        space_before_pt=2.0,
        space_after_pt=4.0,
        line_spacing=1.0,
    )


def add_body_section(document: Document, section_block: SectionBlock, current_columns: int) -> int:
    if section_block.heading and section_block.heading != "正文":
        add_heading_paragraph(document, section_block.heading, level=section_block.level)

    for item in section_block.items:
        if item.type == "paragraph":
            add_paragraph(
                document,
                item.text,
                font_size=BODY_FONT_SIZE_PT,
                first_line_indent_cm=BODY_FIRST_LINE_INDENT_CM,
                line_spacing=BODY_LINE_SPACING,
                space_before_pt=0.0,
                space_after_pt=3.0,
            )
        elif item.type == "figure" and item.figure is not None:
            figure = item.figure
            if figure.is_wide and current_columns != 1:
                single_section = document.add_section(WD_SECTION_START.CONTINUOUS)
                configure_section(single_section, columns=1)
                current_columns = 1
                add_figure_block(document, figure, single_column=True)
                two_col_section = document.add_section(WD_SECTION_START.CONTINUOUS)
                configure_section(two_col_section, columns=2)
                current_columns = 2
            else:
                add_figure_block(document, figure, single_column=(current_columns == 1))
        elif item.type == "table" and item.table is not None:
            table_item = item.table
            if current_columns != 1:
                single_section = document.add_section(WD_SECTION_START.CONTINUOUS)
                configure_section(single_section, columns=1)
                current_columns = 1
                add_table_with_caption(document, table_item, single_column=True)
                two_col_section = document.add_section(WD_SECTION_START.CONTINUOUS)
                configure_section(two_col_section, columns=2)
                current_columns = 2
            else:
                add_table_with_caption(document, table_item, single_column=True)

    return current_columns


def add_body_and_references(document: Document, article: ArticleStructure) -> None:
    body_section = document.add_section(WD_SECTION_START.CONTINUOUS)
    configure_section(body_section, columns=2)
    current_columns = 2

    for section_block in article.sections or build_sections_from_body_items(article.body_items):
        current_columns = add_body_section(document, section_block, current_columns)

    if article.references:
        ref_section = document.add_section(WD_SECTION_START.CONTINUOUS)
        configure_section(ref_section, columns=1)
        add_heading_paragraph(document, "参考文献" if contains_chinese("".join(article.references)) else "References")
        for reference in article.references:
            add_paragraph(
                document,
                reference,
                font_size=10.0,
                first_line_indent_cm=0.0,
                left_indent_cm=0.74,
                hanging_indent_cm=0.74,
                space_before_pt=0.0,
                space_after_pt=2.0,
                line_spacing=1.05,
            )


def create_word_document(article: ArticleStructure, output_path: str) -> None:
    document = Document()
    create_academic_styles(document)
    add_front_matter(document, article)
    add_body_and_references(document, article)
    document.save(output_path)


def build_plain_text_document(article: ArticleStructure) -> str:
    lines = [
        article.title,
        f"Authors: {', '.join(article.authors) if article.authors else 'N/A'}",
        f"Affiliations: {'; '.join(article.affiliations) if article.affiliations else 'N/A'}",
        "",
        "Abstract",
        article.abstract or "N/A",
        "",
        f"Keywords: {'; '.join(article.keywords) if article.keywords else 'N/A'}",
        "",
        "Body",
    ]

    for item in article.body_items:
        if item["type"] == "figure":
            figure = item["figure"]
            lines.extend(["", f"[Figure] {figure.caption}"])
        else:
            lines.extend(["", item["text"]])

    if article.references:
        lines.extend(["", "References"])
        lines.extend(article.references)

    return "\n".join(lines).strip() + "\n"


def build_markdown_document(article: ArticleStructure, paper: Dict[str, Any]) -> str:
    sections = [
        f"# {article.title}",
        "",
        f"- Authors: {', '.join(article.authors) if article.authors else 'N/A'}",
        f"- Affiliations: {'; '.join(article.affiliations) if article.affiliations else 'N/A'}",
        f"- PDF: {paper.get('pdf_url', '')}",
        "",
        "## Abstract",
        "",
        article.abstract or "N/A",
        "",
        f"**Keywords:** {'; '.join(article.keywords) if article.keywords else 'N/A'}",
        "",
        "## Body",
    ]

    for item in article.body_items:
        if item["type"] == "heading":
            sections.extend(["", f"### {item['text']}"])
        elif item["type"] == "paragraph":
            sections.extend(["", item["text"]])
        elif item["type"] == "figure":
            figure = item["figure"]
            sections.extend(["", f"> Figure: {figure.caption}"])

    if article.references:
        sections.extend(["", "## References", ""])
        for ref in article.references:
            sections.append(f"- {ref}")

    return "\n".join(sections).strip() + "\n"


EXPORT_FORMATS = {
    "docx": {
        "extension": ".docx",
        "label": "Word",
        "mimetype": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    },
    "txt": {
        "extension": ".txt",
        "label": "TXT",
        "mimetype": "text/plain; charset=utf-8",
    },
    "md": {
        "extension": ".md",
        "label": "Markdown",
        "mimetype": "text/markdown; charset=utf-8",
    },
}


def write_export_file(export_format: str, paper: Dict[str, Any], article: ArticleStructure, output_path: str) -> None:
    if export_format == "docx":
        create_word_document(article, output_path)
        return

    if export_format == "txt":
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(build_plain_text_document(article))
        return

    if export_format == "md":
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(build_markdown_document(article, paper))
        return

    raise ValueError("不支持的导出格式。")


def process_paper_export(paper: Dict[str, Any], export_format: str = "docx") -> Dict[str, Any]:
    title = paper.get("title", "paper")
    pdf_url = (paper.get("pdf_url") or "").strip()
    export_format = (export_format or "docx").strip().lower()

    if not pdf_url:
        raise ValueError("缺少 PDF 链接。")
    if export_format not in EXPORT_FORMATS:
        raise ValueError("不支持的导出格式。")

    work_dir = tempfile.mkdtemp(prefix="arxiv_word_", dir=tempfile.gettempdir())
    pdf_path = os.path.join(work_dir, "source.pdf")

    safe_title = sanitize_filename(title, default="paper")
    format_info = EXPORT_FORMATS[export_format]
    output_filename = f"{safe_title}{format_info['extension']}"
    output_path = os.path.join(work_dir, output_filename)

    try:
        download_pdf(pdf_url, pdf_path)
        article = parse_pdf_to_article(pdf_path, work_dir, paper)
        if export_format == "docx":
            article = translate_article_to_chinese(article, model=OLLAMA_MODEL)
        write_export_file(export_format, paper, article, output_path)
        safe_remove(pdf_path)

        if not os.path.exists(output_path):
            raise RuntimeError("导出文件生成失败。")

        file_id = uuid.uuid4().hex
        with GENERATED_FILES_LOCK:
            GENERATED_FILES[file_id] = {
                "file_path": output_path,
                "download_name": output_filename,
                "mimetype": format_info["mimetype"],
                "work_dir": work_dir,
                "created_at": time.time(),
            }

        return {
            "file_id": file_id,
            "download_name": output_filename,
            "export_format": export_format,
            "export_label": format_info["label"],
            "page_count": article.page_count,
            "image_count": article.image_count,
        }
    except Exception:
        safe_rmtree(work_dir)
        raise


@app.route("/", methods=["GET"])
def index():
    return render_template("index.html")


@app.route("/search", methods=["POST"])
def search():
    data = request.get_json(silent=True) or request.form
    keyword = (data.get("keyword") or "").strip()

    if not keyword:
        return jsonify({"success": False, "message": "请输入检索关键词。"}), 400

    try:
        results = search_arxiv(keyword, max_results=50)
        return jsonify(
            {
                "success": True,
                "message": f"检索完成，共返回 {len(results)} 篇论文。",
                "results": results,
            }
        )
    except requests.RequestException:
        return jsonify(
            {
                "success": False,
                "message": "访问 arXiv API 失败，请检查网络后重试。",
            }
        ), 502
    except Exception as e:
        app.logger.exception("搜索论文时发生异常：%s", e)
        return jsonify({"success": False, "message": f"检索失败：{e}"}), 500


@app.route("/generate_word", methods=["POST"])
def generate_word():
    data = request.get_json(silent=True) or {}

    title = (data.get("title") or "").strip()
    pdf_url = (data.get("pdf_url") or "").strip()
    if not title or not pdf_url:
        return jsonify({"success": False, "message": "缺少必要的论文信息。"}), 400

    paper = {
        "title": title,
        "authors": data.get("authors", []) or [],
        "year": data.get("year", "未知"),
        "summary": data.get("summary", "") or "",
        "pdf_url": pdf_url,
    }
    export_format = (data.get("export_format") or "docx").strip().lower()

    try:
        result = process_paper_export(paper, export_format=export_format)
        return jsonify(
            {
                "success": True,
                "message": f"{result['export_label']} 文件生成成功。",
                "download_url": url_for("download_file", file_id=result["file_id"]),
                "filename": result["download_name"],
                "export_format": result["export_format"],
                "export_label": result["export_label"],
                "page_count": result["page_count"],
                "image_count": result["image_count"],
            }
        )
    except ValueError as e:
        return jsonify({"success": False, "message": str(e)}), 400
    except RuntimeError as e:
        return jsonify({"success": False, "message": str(e)}), 500
    except Exception as e:
        app.logger.exception("生成 Word 时发生异常：%s", e)
        return jsonify({"success": False, "message": f"生成失败：{e}"}), 500


@app.route("/download/<file_id>", methods=["GET"])
def download_file(file_id: str):
    with GENERATED_FILES_LOCK:
        info = GENERATED_FILES.get(file_id)

    if not info:
        return "文件不存在或已过期。", 404

    file_path = info.get("file_path")
    if not file_path or not os.path.exists(file_path):
        cleanup_generated_file(file_id)
        return "文件不存在或已过期。", 404

    response = send_file(
        file_path,
        as_attachment=True,
        download_name=info.get("download_name", "output.docx"),
        mimetype=info.get("mimetype", "application/octet-stream"),
        max_age=0,
    )
    response.call_on_close(lambda: cleanup_generated_file(file_id))
    return response


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
